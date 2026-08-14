#!/usr/bin/env python3
"""Build the EVOLVE technical report and defense script as polished DOCX files."""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243142"
MUTED = "64748B"
PALE_BLUE = "E8EEF5"
PALE_CYAN = "EAF6F8"
CODE_FILL = "F4F7FA"
RULE = "CBD5E1"
WHITE = "FFFFFF"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_style_font(style, name="Calibri", east_asia="Microsoft YaHei", size=11, color=INK, bold=False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_border(paragraph, bottom=None, left=None):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge, color, size, space in (
        ("bottom", bottom, "12", "5"),
        ("left", left, "18", "8"),
    ):
        if not color:
            continue
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), space)
        tag.set(qn("w:color"), color)
        p_bdr.append(tag)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, result, end])
    set_run_font(run, size=9, color=MUTED)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "打开文档后按 Ctrl+A、F9 可刷新目录页码。"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])


def configure_page(section, title):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    # Pandoc/reference documents may already contain the generated header.
    # Clear it before applying the canonical page furniture so postprocessing
    # remains idempotent.
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"EVOLVE  |  {title}")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    set_paragraph_border(p, bottom=RULE)

    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("EVOLVE  ·  ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_field(p)


def add_or_get_style(doc, name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    set_style_font(title, size=29, color=DARK_BLUE, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, size=14, color=BLUE)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
        "Heading 4": (11, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for name in ("List Bullet", "List Number", "List Bullet 2", "List Number 2"):
        if name in styles:
            style = styles[name]
            set_style_font(style, size=11, color=INK)
            style.paragraph_format.space_after = Pt(4)
            style.paragraph_format.line_spacing = 1.25

    source = add_or_get_style(doc, "Source Code")
    set_style_font(source, name="Consolas", east_asia="Microsoft YaHei", size=8.5, color=INK)
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(7)
    source.paragraph_format.line_spacing = 1.05
    source.paragraph_format.keep_together = False
    source.paragraph_format.left_indent = Inches(0.14)
    source.paragraph_format.right_indent = Inches(0.14)

    quote = styles["Block Text"] if "Block Text" in styles else add_or_get_style(doc, "Block Text")
    set_style_font(quote, size=10.2, color=DARK_BLUE)
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.15)
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(8)
    quote.paragraph_format.line_spacing = 1.15

    caption = styles["Caption"]
    set_style_font(caption, size=9, color=MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True


def create_reference_docx(path: Path, title: str):
    doc = Document()
    configure_page(doc.sections[0], title)
    configure_styles(doc)
    doc.core_properties.title = title
    doc.core_properties.subject = "EVOLVE 操作系统实验教学平台"
    doc.core_properties.author = "EVOLVE Project Team"
    doc.save(path)


def strip_source_title(md_text: str) -> str:
    lines = md_text.splitlines()
    if lines and lines[0].startswith("# "):
        lines = lines[1:]
    while lines and not lines[0].strip():
        lines.pop(0)
    if lines and lines[0].startswith("> Evolving Virtual OS"):
        lines = lines[1:]
        if lines and lines[0].startswith("> 版本基线"):
            lines = lines[1:]
    while lines and (not lines[0].strip() or lines[0].strip() == "---"):
        lines.pop(0)
    return "\n".join(lines).strip() + "\n"


def raw_page_break():
    return '```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```'


def raw_toc():
    return """```{=openxml}
<w:p>
  <w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>
  <w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>
  <w:r><w:fldChar w:fldCharType="separate"/></w:r>
  <w:r><w:t>打开文档后按 Ctrl+A、F9 可刷新目录页码。</w:t></w:r>
  <w:r><w:fldChar w:fldCharType="end"/></w:r>
</w:p>
```"""


def prepare_markdown(source: Path, output: Path, kind: str):
    body = strip_source_title(source.read_text(encoding="utf-8"))
    if kind == "report":
        title = "EVOLVE 总体实验技术文档"
        subtitle = "Evolving Virtual OS Learning & Verification Environment"
        kicker = "总体实验报告 · 技术实现与代码闭合审查"
        date = "版本基线：2026 年 8 月 14 日"
    else:
        title = "EVOLVE 答辩演讲稿"
        subtitle = "配套《EVOLVE 答辩演示》逐页讲稿"
        kicker = "技术方案深层讲解版"
        date = "演讲基线：2026 年 8 月 14 日"
    front = f"""---
title: "{title}"
subtitle: "{subtitle}"
author: "{kicker}"
date: "{date}"
---

{raw_page_break()}

# 目录

{raw_toc()}

{raw_page_break()}

"""
    output.write_text(front + body, encoding="utf-8")


def content_weight(text):
    chinese = len(re.findall(r"[\u4e00-\u9fff]", text))
    ascii_chars = len(text) - chinese
    return max(4, chinese * 2 + ascii_chars)


def table_widths(table, total=9360):
    cols = len(table.columns)
    weights = []
    for idx in range(cols):
        texts = [row.cells[idx].text.strip() for row in table.rows]
        samples = sorted((content_weight(t) for t in texts), reverse=True)[:4]
        weights.append(max(8, sum(samples) / max(1, len(samples))))
    minimums = [760 if cols >= 4 else 1000] * cols
    raw_total = sum(weights)
    allocated = [max(minimums[i], int(total * weights[i] / raw_total)) for i in range(cols)]
    while sum(allocated) > total:
        reducible = [i for i, value in enumerate(allocated) if value > minimums[i]]
        if not reducible:
            break
        i = max(reducible, key=lambda n: allocated[n] - minimums[n])
        allocated[i] -= min(20, allocated[i] - minimums[i])
    while sum(allocated) < total:
        i = max(range(cols), key=lambda n: weights[n])
        allocated[i] += min(20, total - sum(allocated))
    return allocated


def style_tables(doc):
    for table in doc.tables:
        table.autofit = False
        table.alignment = 0
        table_pr = table._tbl.tblPr
        tbl_w = table_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            table_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), "9360")
        tbl_w.set(qn("w:type"), "dxa")
        tbl_ind = table_pr.find(qn("w:tblInd"))
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            table_pr.append(tbl_ind)
        tbl_ind.set(qn("w:w"), "120")
        tbl_ind.set(qn("w:type"), "dxa")

        widths = table_widths(table)
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)

        if table.rows:
            set_repeat_table_header(table.rows[0])
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                set_cell_width(cell, widths[col_idx])
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_idx == 0:
                    shade_cell(cell, PALE_BLUE)
                elif row_idx % 2 == 0:
                    shade_cell(cell, "F8FAFC")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2.5)
                    paragraph.paragraph_format.line_spacing = 1.08
                    for run in paragraph.runs:
                        set_run_font(run, size=8.8 if len(table.columns) >= 4 else 9.2, color=INK)
                        if row_idx == 0:
                            run.bold = True
                            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)


def style_paragraphs(doc):
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()
        if style_name == "Source Code":
            p_pr = paragraph._p.get_or_add_pPr()
            shd = p_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                p_pr.append(shd)
            shd.set(qn("w:fill"), CODE_FILL)
            set_paragraph_border(paragraph, left=BLUE)
            for run in paragraph.runs:
                set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=8.5, color=INK)
        elif style_name == "Block Text":
            p_pr = paragraph._p.get_or_add_pPr()
            shd = p_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                p_pr.append(shd)
            shd.set(qn("w:fill"), PALE_CYAN)
            set_paragraph_border(paragraph, left=BLUE)
        elif style_name == "Title":
            paragraph.paragraph_format.space_before = Pt(136)
        elif style_name == "Subtitle":
            paragraph.paragraph_format.space_after = Pt(16)
        elif style_name == "Author":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(56)
            paragraph.paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                set_run_font(run, size=10.5, color=MUTED, bold=True)
        elif style_name == "Date":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=10, color=MUTED)
        elif style_name == "Heading 1":
            set_paragraph_border(paragraph, bottom=RULE)
        if text == "目录":
            paragraph.paragraph_format.space_before = Pt(16)
            paragraph.paragraph_format.space_after = Pt(14)


def keep_table_context(doc):
    for table in doc.tables:
        previous = table._tbl.getprevious()
        if previous is not None and previous.tag == qn("w:p"):
            p_pr = previous.get_or_add_pPr()
            keep = p_pr.find(qn("w:keepNext"))
            if keep is None:
                p_pr.append(OxmlElement("w:keepNext"))


def update_fields_on_open(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def postprocess_docx(path: Path, title: str):
    doc = Document(path)
    for section in doc.sections:
        configure_page(section, title)
    configure_styles(doc)
    style_paragraphs(doc)
    style_tables(doc)
    keep_table_context(doc)
    update_fields_on_open(doc)
    doc.core_properties.title = title
    doc.core_properties.subject = "EVOLVE 操作系统实验教学平台"
    doc.core_properties.author = "EVOLVE Project Team"
    doc.save(path)


def build(source: Path, output: Path, kind: str):
    title = "总体实验技术文档" if kind == "report" else "答辩演讲稿"
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="evolve-doc-") as temp_dir:
        temp = Path(temp_dir)
        reference = temp / "reference.docx"
        prepared = temp / "prepared.md"
        create_reference_docx(reference, title)
        prepare_markdown(source, prepared, kind)
        cmd = [
            shutil.which("pandoc") or "pandoc",
            str(prepared),
            "--from=gfm+raw_attribute",
            "--to=docx",
            f"--reference-doc={reference}",
            "--standalone",
            f"--output={output}",
        ]
        subprocess.run(cmd, check=True)
    postprocess_docx(output, title)
    print(f"built: {output}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--kind", choices=("report", "speech"), required=True)
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve(), args.kind)


if __name__ == "__main__":
    main()
